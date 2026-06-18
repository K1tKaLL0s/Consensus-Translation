from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path


@dataclass(frozen=True)
class AgentInputDocument:
    input_ref: str
    text: str


def _load_docx(path: Path) -> str:
    from docx import Document

    document = Document(path)
    paragraphs = [
        paragraph.text.strip()
        for paragraph in document.paragraphs
        if paragraph.text.strip()
    ]
    return "\n".join(paragraphs)


def load_agent_input(path: str | Path) -> AgentInputDocument:
    source = Path(path)
    suffix = source.suffix.lower()
    if suffix in {".txt", ".md"}:
        text = source.read_text(encoding="utf-8")
    elif suffix == ".docx":
        text = _load_docx(source)
    else:
        raise ValueError(f"unsupported input file type: {suffix}")
    return AgentInputDocument(input_ref=str(source), text=text)


def load_batch_inputs(paths: list[str | Path]) -> list[AgentInputDocument]:
    return [load_agent_input(path) for path in paths]
