from pathlib import Path

from docx import Document


class FileIngestService:
    SUPPORTED_SUFFIXES = {".txt", ".md", ".docx"}

    def validate_suffix(self, suffix: str) -> str:
        normalized = suffix.strip().lower()
        if not normalized.startswith("."):
            normalized = f".{normalized}"
        if normalized not in self.SUPPORTED_SUFFIXES:
            raise ValueError("unsupported file type")
        return normalized

    def validate_path(self, path: Path) -> str:
        return self.validate_suffix(path.suffix)

    def read_text(self, path: Path) -> str:
        suffix = self.validate_path(path)
        if suffix in {".txt", ".md"}:
            return path.read_text(encoding="utf-8")

        document = Document(path)
        return "\n".join(paragraph.text for paragraph in document.paragraphs)

    def write_text(self, path: Path, text: str) -> None:
        suffix = self.validate_path(path)
        path.parent.mkdir(parents=True, exist_ok=True)

        if suffix in {".txt", ".md"}:
            path.write_text(text, encoding="utf-8")
            return

        document = Document()
        lines = text.splitlines() or [text]
        for line in lines:
            document.add_paragraph(line)
        document.save(path)
