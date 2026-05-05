from pathlib import Path

import pytest
from docx import Document

from src.services.file_ingest_service import FileIngestService


def test_read_txt_file(tmp_path: Path) -> None:
    path = tmp_path / "sample.txt"
    path.write_text("hello", encoding="utf-8")

    text = FileIngestService().read_text(path)

    assert text == "hello"


def test_read_docx_file(tmp_path: Path) -> None:
    path = tmp_path / "sample.docx"
    document = Document()
    document.add_paragraph("line1")
    document.add_paragraph("line2")
    document.save(path)

    text = FileIngestService().read_text(path)

    assert "line1" in text
    assert "line2" in text


def test_write_md_file(tmp_path: Path) -> None:
    path = tmp_path / "output.md"

    FileIngestService().write_text(path, "# title")

    assert path.read_text(encoding="utf-8") == "# title"


def test_validate_suffix_rejects_pdf() -> None:
    with pytest.raises(ValueError, match="unsupported file type"):
        FileIngestService().validate_suffix(".pdf")
