from pathlib import Path
import sys

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "src"
if str(SRC) not in sys.path:
    sys.path.insert(0, str(SRC))


from consensus_translation.agent_inputs import load_agent_input, load_batch_inputs


def test_load_agent_input_supports_txt_md_and_docx(tmp_path):
    txt_path = tmp_path / "sample.txt"
    md_path = tmp_path / "notes.md"
    docx_path = tmp_path / "chapter.docx"

    txt_path.write_text("plain text", encoding="utf-8")
    md_path.write_text("# Title\n\nmarkdown body", encoding="utf-8")
    doc = Document()
    doc.add_paragraph("第一段")
    doc.add_paragraph("第二段")
    doc.save(docx_path)

    assert load_agent_input(txt_path).text == "plain text"
    assert load_agent_input(md_path).text == "# Title\n\nmarkdown body"
    assert load_agent_input(docx_path).text == "第一段\n第二段"


def test_load_batch_inputs_preserves_order_and_refs(tmp_path):
    first = tmp_path / "a.txt"
    second = tmp_path / "b.md"
    first.write_text("A", encoding="utf-8")
    second.write_text("B", encoding="utf-8")

    docs = load_batch_inputs([first, second])

    assert [doc.text for doc in docs] == ["A", "B"]
    assert [doc.input_ref for doc in docs] == [str(first), str(second)]
