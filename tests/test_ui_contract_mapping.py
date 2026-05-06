from pathlib import Path
import sys


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "src"
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))
if str(SRC) not in sys.path:
    sys.path.insert(0, str(SRC))


from app import (
    PAGE_FIELD_MAP,
    PAGE_LABEL_MAP,
    extract_uploaded_text,
    extract_page_data,
    get_page_select_keys,
    resolve_input_text,
    resolve_dot_path,
)


def test_monitor_page_uses_plan_defined_status_fields():
    assert PAGE_FIELD_MAP["monitor"] == [
        "stage_status.current",
        "stage_status.progress",
        "stage_status.retry_count",
        "stage_status.error_code",
        "stage_status.error_message",
    ]


def test_mdwc_page_contains_plan_defined_subset():
    mdwc_fields = set(PAGE_FIELD_MAP["mdwc"])

    assert {
        "weights",
        "token_score",
        "sentence_score",
        "segment_score",
        "user_prior",
        "final_score",
        "decision_reason",
    }.issubset(mdwc_fields)


def test_resolve_dot_path_returns_nested_value_and_none_for_missing():
    payload = {
        "stage_status": {
            "current": "finalize",
            "progress": 1.0,
        }
    }

    assert resolve_dot_path(payload, "stage_status.current") == "finalize"
    assert resolve_dot_path(payload, "stage_status.progress") == 1.0
    assert resolve_dot_path(payload, "stage_status.error_code") is None


def test_extract_page_data_maps_runtime_payload_values():
    payload = {
        "contract": {
            "job_id": "job-123",
            "mode": "local",
            "source_lang": "zh",
            "target_lang": "ja",
            "topic": "travel",
            "stage_status": {
                "current": "review",
                "progress": 0.8,
                "retry_count": 1,
                "error_code": None,
                "error_message": None,
            },
        },
        "weights": {"token": 0.4},
        "token_score": 0.45,
        "sentence_score": 0.45,
        "segment_score": 0.45,
        "user_prior": 0.5,
        "final_score": 0.4525,
        "decision_reason": "left-score-greater-or-equal",
    }

    config = extract_page_data("config", payload)
    monitor = extract_page_data("monitor", payload)
    mdwc = extract_page_data("mdwc", payload)

    assert config["job_id"] == "job-123"
    assert config["source_lang"] == "zh"
    assert config["target_lang"] == "ja"
    assert config["topic"] == "travel"
    assert monitor["stage_status.current"] == "review"
    assert monitor["stage_status.progress"] == 0.8
    assert mdwc["token_score"] == 0.45
    assert mdwc["decision_reason"] == "left-score-greater-or-equal"


def test_monitor_fields_fallback_to_contract_when_runtime_key_missing():
    payload = {
        "contract": {
            "stage_status": {
                "current": "finalize",
                "progress": 1.0,
                "retry_count": 0,
                "error_code": None,
                "error_message": None,
            }
        },
        "final_score": 0.9,
    }

    data = extract_page_data("monitor", payload)

    assert data["stage_status.current"] == "finalize"
    assert data["stage_status.progress"] == 1.0


def test_monitor_field_preserves_runtime_none_without_contract_fallback():
    payload = {
        "stage_status": {
            "current": None,
        },
        "contract": {
            "stage_status": {
                "current": "finalize",
                "progress": 1.0,
                "retry_count": 0,
                "error_code": None,
                "error_message": None,
            }
        },
    }

    data = extract_page_data("monitor", payload)

    assert data["stage_status.current"] is None


def test_ui_does_not_expose_phase3_ai_mode_controls():
    forbidden = {"ai_mode", "ai_vote", "ai_iteration", "multi_model"}
    all_fields = {field for fields in PAGE_FIELD_MAP.values() for field in fields}

    assert forbidden.isdisjoint(all_fields)


def test_page_field_map_keys_stay_stable_for_backend_contract():
    assert list(PAGE_FIELD_MAP.keys()) == [
        "config",
        "monitor",
        "compare",
        "mdwc",
        "revision",
        "pretrain_report",
    ]


def test_page_label_map_uses_chinese_labels_for_display_only():
    assert PAGE_LABEL_MAP == {
        "config": "配置",
        "monitor": "监控",
        "compare": "对比",
        "mdwc": "MDWC评分",
        "revision": "修订",
        "pretrain_report": "预训练报告",
    }
    assert set(PAGE_LABEL_MAP.keys()) == set(PAGE_FIELD_MAP.keys())


def test_page_label_values_are_unique_to_avoid_display_collisions():
    labels = list(PAGE_LABEL_MAP.values())

    assert len(labels) == len(set(labels))


def test_page_selector_options_use_stable_contract_keys():
    assert get_page_select_keys() == list(PAGE_FIELD_MAP.keys())


class DummyUpload:
    def __init__(self, name: str, file_type: str, data: bytes) -> None:
        self.name = name
        self.type = file_type
        self._data = data

    def getvalue(self) -> bytes:
        return self._data


def test_extract_uploaded_text_supports_utf8_plain_text_and_md():
    txt_upload = DummyUpload("note.txt", "text/plain", "你好 world".encode("utf-8"))
    md_upload = DummyUpload("note.md", "text/markdown", "# 标题\n内容".encode("utf-8"))

    txt_text, txt_meta = extract_uploaded_text(txt_upload)
    md_text, md_meta = extract_uploaded_text(md_upload)

    assert txt_text == "你好 world"
    assert txt_meta["ok"] is True
    assert txt_meta["file_type"] == "text/plain"
    assert md_text == "# 标题\n内容"
    assert md_meta["ok"] is True
    assert md_meta["file_type"] == "text/markdown"


def test_extract_uploaded_text_supports_docx_from_bytes():
    from io import BytesIO

    from docx import Document

    buf = BytesIO()
    doc = Document()
    doc.add_paragraph("第一段")
    doc.add_paragraph("Second line")
    doc.save(buf)

    upload = DummyUpload(
        "sample.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        buf.getvalue(),
    )

    text, meta = extract_uploaded_text(upload)

    assert text == "第一段\nSecond line"
    assert meta["ok"] is True
    assert meta["file_ext"] == ".docx"


def test_extract_uploaded_text_marks_unsupported_type_as_non_ok():
    upload = DummyUpload("sample.pdf", "application/pdf", b"%PDF")

    text, meta = extract_uploaded_text(upload)

    assert text == ""
    assert meta["ok"] is False
    assert meta["reason"] == "unsupported_type"


def test_extract_uploaded_text_returns_decode_error_for_invalid_utf8_text():
    upload = DummyUpload("broken.txt", "text/plain", b"\xff\xfe\xfa")

    text, meta = extract_uploaded_text(upload)

    assert text == ""
    assert meta["ok"] is False
    assert meta["reason"] == "decode_error"


def test_extract_uploaded_text_returns_docx_parse_error_for_malformed_docx_bytes():
    upload = DummyUpload(
        "broken.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        b"not-a-zip-docx",
    )

    text, meta = extract_uploaded_text(upload)

    assert text == ""
    assert meta["ok"] is False
    assert meta["reason"] == "docx_parse_error"


def test_extract_uploaded_text_returns_docx_dependency_missing_when_docx_unavailable(monkeypatch):
    import builtins

    original_import = builtins.__import__

    def fake_import(name, globals=None, locals=None, fromlist=(), level=0):
        if name == "docx":
            raise ModuleNotFoundError("No module named 'docx'")
        return original_import(name, globals, locals, fromlist, level)

    monkeypatch.setattr(builtins, "__import__", fake_import)
    upload = DummyUpload(
        "sample.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        b"docx-bytes",
    )

    text, meta = extract_uploaded_text(upload)

    assert text == ""
    assert meta["ok"] is False
    assert meta["reason"] == "docx_dependency_missing"


def test_resolve_input_text_prefers_uploaded_non_empty_text():
    chosen_text, chosen_meta = resolve_input_text("手动输入", "上传文本", {"ok": True})

    assert chosen_text == "上传文本"
    assert chosen_meta["source"] == "upload"


def test_resolve_input_text_falls_back_to_manual_when_upload_empty_or_not_ok():
    text_empty, meta_empty = resolve_input_text("手动输入", "   ", {"ok": True})
    text_bad, meta_bad = resolve_input_text("手动输入", "上传文本", {"ok": False})

    assert text_empty == "手动输入"
    assert meta_empty["source"] == "manual"
    assert text_bad == "手动输入"
    assert meta_bad["source"] == "manual"
